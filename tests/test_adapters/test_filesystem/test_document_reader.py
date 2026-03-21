"""Tests for MultiFormatDocumentReader adapter."""

import warnings
from pathlib import Path

import pytest

from mkcv.adapters.filesystem.document_reader import MultiFormatDocumentReader
from mkcv.core.exceptions.kb_generation import DocumentReadError
from mkcv.core.models.document_content import DocumentContent


@pytest.fixture
def reader() -> MultiFormatDocumentReader:
    """Create a fresh reader instance."""
    return MultiFormatDocumentReader()


# ------------------------------------------------------------------
# supported_extensions
# ------------------------------------------------------------------


class TestSupportedExtensions:
    """Tests for supported_extensions()."""

    def test_returns_set(self, reader: MultiFormatDocumentReader) -> None:
        exts = reader.supported_extensions()
        assert isinstance(exts, set)

    def test_includes_pdf(self, reader: MultiFormatDocumentReader) -> None:
        assert ".pdf" in reader.supported_extensions()

    def test_includes_markdown(self, reader: MultiFormatDocumentReader) -> None:
        exts = reader.supported_extensions()
        assert ".md" in exts
        assert ".markdown" in exts

    def test_includes_text(self, reader: MultiFormatDocumentReader) -> None:
        exts = reader.supported_extensions()
        assert ".txt" in exts
        assert ".text" in exts

    def test_includes_docx(self, reader: MultiFormatDocumentReader) -> None:
        assert ".docx" in reader.supported_extensions()

    def test_includes_html(self, reader: MultiFormatDocumentReader) -> None:
        exts = reader.supported_extensions()
        assert ".html" in exts
        assert ".htm" in exts


# ------------------------------------------------------------------
# read_file: text files
# ------------------------------------------------------------------


class TestReadText:
    """Tests for reading plain text files."""

    def test_reads_txt_file(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        f = tmp_path / "notes.txt"
        f.write_text("Hello world", encoding="utf-8")
        result = reader.read_file(f)
        assert result.text == "Hello world"
        assert result.format == "text"
        assert result.char_count == 11

    def test_reads_text_extension(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        f = tmp_path / "notes.text"
        f.write_text("Alt extension", encoding="utf-8")
        result = reader.read_file(f)
        assert result.format == "text"
        assert result.text == "Alt extension"


# ------------------------------------------------------------------
# read_file: markdown files
# ------------------------------------------------------------------


class TestReadMarkdown:
    """Tests for reading Markdown files."""

    def test_reads_md_file(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        f = tmp_path / "resume.md"
        f.write_text("# Resume\n\nJane Doe", encoding="utf-8")
        result = reader.read_file(f)
        assert result.format == "markdown"
        assert "Jane Doe" in result.text

    def test_reads_markdown_extension(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        f = tmp_path / "resume.markdown"
        f.write_text("# Resume", encoding="utf-8")
        result = reader.read_file(f)
        assert result.format == "markdown"


# ------------------------------------------------------------------
# read_file: HTML files
# ------------------------------------------------------------------


class TestReadHTML:
    """Tests for reading HTML files."""

    def test_reads_html_file(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        html = (
            "<html><head><title>My CV</title></head><body><p>Engineer</p></body></html>"
        )
        f = tmp_path / "profile.html"
        f.write_text(html, encoding="utf-8")
        result = reader.read_file(f)
        assert result.format == "html"
        assert "Engineer" in result.text

    def test_html_extracts_title_metadata(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        html = "<html><head><title>My CV</title></head><body>Text</body></html>"
        f = tmp_path / "profile.html"
        f.write_text(html, encoding="utf-8")
        result = reader.read_file(f)
        assert result.metadata.get("title") == "My CV"

    def test_reads_htm_extension(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        f = tmp_path / "page.htm"
        f.write_text("<html><body>Hello</body></html>", encoding="utf-8")
        result = reader.read_file(f)
        assert result.format == "html"


# ------------------------------------------------------------------
# read_file: DOCX files
# ------------------------------------------------------------------


class TestReadDocx:
    """Tests for reading DOCX files."""

    def test_reads_docx_file(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        """Test reading a real DOCX file created with python-docx."""
        import docx as python_docx

        f = tmp_path / "resume.docx"
        doc = python_docx.Document()
        doc.core_properties.title = "My Resume"
        doc.core_properties.author = "Test Author"
        doc.add_paragraph("Senior Software Engineer")
        doc.add_paragraph("Built scalable systems")
        doc.save(str(f))

        result = reader.read_file(f)
        assert result.format == "docx"
        assert "Senior Software Engineer" in result.text
        assert result.metadata.get("title") == "My Resume"
        assert result.metadata.get("author") == "Test Author"


# ------------------------------------------------------------------
# read_file: PDF files
# ------------------------------------------------------------------


class TestReadPDF:
    """Tests for reading PDF files."""

    def test_reads_pdf_file(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        """Test reading a real PDF file created with pypdf."""
        from pypdf import PdfWriter

        f = tmp_path / "resume.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        with open(f, "wb") as fp:
            writer.write(fp)

        # This PDF has no text, so it should trigger the image-only warning
        with warnings.catch_warnings(record=True) as w:
            warnings.simplefilter("always")
            result = reader.read_file(f)
            assert any("image-only" in str(warning.message) for warning in w)

        assert result.format == "pdf"

    def test_corrupted_pdf_raises_document_read_error(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        f = tmp_path / "bad.pdf"
        f.write_bytes(b"not a pdf file at all")
        with pytest.raises(DocumentReadError, match="Cannot read PDF"):
            reader.read_file(f)


# ------------------------------------------------------------------
# read_file: unsupported formats
# ------------------------------------------------------------------


class TestReadUnsupported:
    """Tests for unsupported file formats."""

    def test_unsupported_extension_raises(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        f = tmp_path / "data.csv"
        f.write_text("a,b,c", encoding="utf-8")
        with pytest.raises(DocumentReadError, match="Unsupported file format"):
            reader.read_file(f)

    def test_nonexistent_file_raises(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        f = tmp_path / "ghost.txt"
        with pytest.raises(DocumentReadError, match="File not found"):
            reader.read_file(f)


# ------------------------------------------------------------------
# read_file: returns DocumentContent
# ------------------------------------------------------------------


class TestReadFileReturnType:
    """Tests that read_file returns properly structured DocumentContent."""

    def test_returns_document_content(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        f = tmp_path / "test.txt"
        f.write_text("content", encoding="utf-8")
        result = reader.read_file(f)
        assert isinstance(result, DocumentContent)

    def test_source_path_is_set(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        f = tmp_path / "test.txt"
        f.write_text("content", encoding="utf-8")
        result = reader.read_file(f)
        assert result.source_path == f


# ------------------------------------------------------------------
# read_sources: files, directories, and globs
# ------------------------------------------------------------------


class TestReadSources:
    """Tests for read_sources() directory scanning."""

    def test_reads_single_file(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        f = tmp_path / "test.txt"
        f.write_text("hello", encoding="utf-8")
        results = reader.read_sources([f])
        assert len(results) == 1
        assert results[0].text == "hello"

    def test_reads_multiple_files(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        f1 = tmp_path / "a.txt"
        f2 = tmp_path / "b.md"
        f1.write_text("file a", encoding="utf-8")
        f2.write_text("file b", encoding="utf-8")
        results = reader.read_sources([f1, f2])
        assert len(results) == 2

    def test_reads_directory(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        subdir = tmp_path / "docs"
        subdir.mkdir()
        (subdir / "a.txt").write_text("aaa", encoding="utf-8")
        (subdir / "b.md").write_text("bbb", encoding="utf-8")
        results = reader.read_sources([subdir])
        assert len(results) == 2

    def test_directory_with_glob_filter(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        subdir = tmp_path / "docs"
        subdir.mkdir()
        (subdir / "a.txt").write_text("aaa", encoding="utf-8")
        (subdir / "b.md").write_text("bbb", encoding="utf-8")
        results = reader.read_sources([subdir], glob="**/*.md")
        assert len(results) == 1
        assert results[0].format == "markdown"

    def test_skips_unsupported_files_in_directory(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        subdir = tmp_path / "docs"
        subdir.mkdir()
        (subdir / "a.txt").write_text("aaa", encoding="utf-8")
        (subdir / "b.csv").write_text("x,y,z", encoding="utf-8")
        results = reader.read_sources([subdir])
        assert len(results) == 1

    def test_skips_unsupported_standalone_files(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        f = tmp_path / "data.csv"
        f.write_text("x,y,z", encoding="utf-8")
        results = reader.read_sources([f])
        assert len(results) == 0

    def test_skips_nonexistent_paths(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        ghost = tmp_path / "nonexistent"
        results = reader.read_sources([ghost])
        assert len(results) == 0

    def test_results_sorted_by_path(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        (tmp_path / "c.txt").write_text("c", encoding="utf-8")
        (tmp_path / "a.txt").write_text("a", encoding="utf-8")
        (tmp_path / "b.txt").write_text("b", encoding="utf-8")
        results = reader.read_sources([tmp_path])
        names = [r.source_path.name for r in results]
        assert names == sorted(names)

    def test_deduplicates_files(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        """When the same file is passed both directly and via directory."""
        f = tmp_path / "test.txt"
        f.write_text("hello", encoding="utf-8")
        results = reader.read_sources([f, tmp_path])
        assert len(results) == 1

    def test_recursive_directory_scan(
        self, reader: MultiFormatDocumentReader, tmp_path: Path
    ) -> None:
        nested = tmp_path / "a" / "b"
        nested.mkdir(parents=True)
        (nested / "deep.txt").write_text("deep content", encoding="utf-8")
        results = reader.read_sources([tmp_path])
        assert len(results) == 1
        assert results[0].text == "deep content"
